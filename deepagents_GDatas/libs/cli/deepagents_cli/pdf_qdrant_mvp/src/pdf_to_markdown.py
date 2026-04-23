"""
文档转 Markdown 转换脚本

将 PDF / DOCX / DOC 文件转换为 Markdown 格式并保存到指定文件夹。
- PDF：使用 pdfplumber 的字符级提取功能，精确保留下标/上标格式。
- DOCX：使用 python-docx 提取段落、表格、标题、列表等结构化内容。
- DOC：先转换为 DOCX（依赖 LibreOffice 或 MS Word），再按 DOCX 流程处理。

使用方法:
    python src/pdf_to_markdown.py --pdf-dir <PDF目录> --output-dir <输出目录>
    python src/pdf_to_markdown.py --pdf-file <PDF文件路径> --output-dir <输出目录>
    python src/pdf_to_markdown.py --docx-file <DOCX文件路径> --output-dir <输出目录>
    python src/pdf_to_markdown.py --doc-file <DOC文件路径> --output-dir <输出目录>
    python src/pdf_to_markdown.py --input-dir <目录> --output-dir <输出目录>
"""

import os
import re
import sys
import argparse
import subprocess
import tempfile
import shutil
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, List, Tuple, Optional
from collections import Counter
from dataclasses import dataclass, field

# 添加父目录到路径
sys.path.insert(0, str(Path(__file__).parent))

from rich.console import Console
from rich.progress import Progress
from rich.panel import Panel

from qdrant_config import config

# 初始化控制台
console = Console()


# ========== 辅助函数 ==========

def get_pdf_files(directory: str) -> List[Dict[str, str]]:
    """获取目录中的所有 PDF 文件
    
    Args:
        directory: 目录路径
        
    Returns:
        包含 PDF 文件信息的列表，每个元素包含:
        - name: 文件名
        - path: 完整路径
        - collection_name: 集合名称（用于向量数据库）
    """
    import re
    from pathlib import Path
    
    pdf_files = []
    dir_path = Path(directory)
    
    if not dir_path.exists():
        return pdf_files
    
    for pdf_file in sorted(dir_path.glob("*.pdf")):
        # 生成集合名称：将文件名转换为小写，替换特殊字符
        name = pdf_file.stem
        collection_name = re.sub(r'[^a-zA-Z0-9_]', '_', name.lower())
        collection_name = re.sub(r'_+', '_', collection_name).strip('_')
        
        pdf_files.append({
            "name": pdf_file.name,
            "path": str(pdf_file),
            "collection_name": collection_name
        })
    
    return pdf_files


# ========== 获取 DOCX / DOC 文件列表 ==========

def get_docx_files(directory: str) -> List[Dict[str, str]]:
    """获取目录中的所有 DOCX 文件

    Args:
        directory: 目录路径

    Returns:
        包含 DOCX 文件信息的列表
    """
    docx_files = []
    dir_path = Path(directory)

    if not dir_path.exists():
        return docx_files

    for docx_file in sorted(dir_path.glob("*.docx")):
        name = docx_file.stem
        collection_name = re.sub(r'[^a-zA-Z0-9_]', '_', name.lower())
        collection_name = re.sub(r'_+', '_', collection_name).strip('_')

        docx_files.append({
            "name": docx_file.name,
            "path": str(docx_file),
            "collection_name": collection_name
        })

    return docx_files


def get_doc_files(directory: str) -> List[Dict[str, str]]:
    """获取目录中的所有 DOC 文件（不含 .docx）

    Args:
        directory: 目录路径

    Returns:
        包含 DOC 文件信息的列表
    """
    doc_files = []
    dir_path = Path(directory)

    if not dir_path.exists():
        return doc_files

    for doc_file in sorted(dir_path.glob("*.doc")):
        # 排除 .docx 文件
        if doc_file.suffix.lower() != '.doc':
            continue
        name = doc_file.stem
        collection_name = re.sub(r'[^a-zA-Z0-9_]', '_', name.lower())
        collection_name = re.sub(r'_+', '_', collection_name).strip('_')

        doc_files.append({
            "name": doc_file.name,
            "path": str(doc_file),
            "collection_name": collection_name
        })

    return doc_files


def get_all_document_files(directory: str) -> List[Dict[str, str]]:
    """获取目录中所有支持的文档文件（PDF + DOCX + DOC）

    Args:
        directory: 目录路径

    Returns:
        包含所有文档文件信息的列表，每个元素额外包含 file_type 字段
    """
    all_files = []

    for f in get_pdf_files(directory):
        f["file_type"] = "pdf"
        all_files.append(f)

    for f in get_docx_files(directory):
        f["file_type"] = "docx"
        all_files.append(f)

    for f in get_doc_files(directory):
        f["file_type"] = "doc"
        all_files.append(f)

    # 按文件名排序
    all_files.sort(key=lambda x: x["name"])

    return all_files


# ========== 常量定义 ==========

@dataclass
class ScriptMaps:
    """下标和上标字符映射"""
    subscript: Dict[str, str] = field(default_factory=lambda: {
        '0': '₀', '1': '₁', '2': '₂', '3': '₃', '4': '₄',
        '5': '₅', '6': '₆', '7': '₇', '8': '₈', '9': '₉',
        '+': '₊', '-': '₋', '=': '₌', '(': '₍', ')': '₎',
        'a': 'ₐ', 'e': 'ₑ', 'h': 'ₕ', 'i': 'ᵢ', 'j': 'ⱼ',
        'k': 'ₖ', 'l': 'ₗ', 'm': 'ₘ', 'n': 'ₙ', 'o': 'ₒ',
        'p': 'ₚ', 'r': 'ᵣ', 's': 'ₛ', 't': 'ₜ', 'u': 'ᵤ',
        'v': 'ᵥ', 'x': 'ₓ',
    })
    superscript: Dict[str, str] = field(default_factory=lambda: {
        '0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴',
        '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹',
        '+': '⁺', '-': '⁻', '=': '⁼', '(': '⁽', ')': '⁾',
        'n': 'ⁿ', 'i': 'ⁱ',
    })


SCRIPT_MAPS = ScriptMaps()

# 化学元素符号集合
CHEMICAL_ELEMENTS = frozenset({
    'H', 'He', 'Li', 'Be', 'B', 'C', 'N', 'O', 'F', 'Ne',
    'Na', 'Mg', 'Al', 'Si', 'P', 'S', 'Cl', 'Ar', 'K', 'Ca',
    'Sc', 'Ti', 'V', 'Cr', 'Mn', 'Fe', 'Co', 'Ni', 'Cu', 'Zn',
    'Ga', 'Ge', 'As', 'Se', 'Br', 'Kr', 'Rb', 'Sr', 'Y', 'Zr',
    'Nb', 'Mo', 'Tc', 'Ru', 'Rh', 'Pd', 'Ag', 'Cd', 'In', 'Sn',
    'Sb', 'Te', 'I', 'Xe', 'Cs', 'Ba', 'La', 'Ce', 'Pr', 'Nd',
    'Pm', 'Sm', 'Eu', 'Gd', 'Tb', 'Dy', 'Ho', 'Er', 'Tm', 'Yb',
    'Lu', 'Hf', 'Ta', 'W', 'Re', 'Os', 'Ir', 'Pt', 'Au', 'Hg',
    'Tl', 'Pb', 'Bi', 'Po', 'At', 'Rn', 'Fr', 'Ra', 'Ac', 'Th',
    'Pa', 'U', 'Np', 'Pu', 'Am', 'Cm', 'Bk', 'Cf', 'Es', 'Fm'
})

# 面积/体积单位（后面跟数字通常是上标，如 cm², mm³）
AREA_VOLUME_UNITS = frozenset({
    'cm', 'mm', 'm', 'km', 'μm', 'nm', 'pm', 'Å'
})

# 时间/速率单位（后面跟负数是上标，如 min⁻¹, s⁻¹）
RATE_UNITS = frozenset({
    'min', 'h', 's', 'ms', 'rpm'
})

# 其他常见单位
OTHER_UNITS = frozenset({
    'mol', 'M', 'mM', 'μM', 'L', 'mL', 'μL', 'C', 'K', 'J', 'kJ',
    'eV', 'V', 'mV', 'A', 'W', 'Hz', 'kHz', 'MHz', 'GHz', 'Pa', 'kPa',
    'wt', 'vol', 'g', 'kg', 'mg'
})

# 所有单位集合
ALL_UNITS = AREA_VOLUME_UNITS | RATE_UNITS | OTHER_UNITS

# 特殊字符映射
SPECIAL_CHAR_MAP = {
    '': '*',
    '−': '-',
    '–': '-',
    '—': '--',
    ''': "'",
    ''': "'",
    '"': '"',
    '"': '"',
    '…': '...',
    '': '×',
    '⨉': '×',
    '✕': '×',
    '': 'μ',
    'µ': 'μ',
}

# 预编译正则表达式（性能优化）
RE_ELEMENT_NUMBER = re.compile(r'\b([A-Z][a-z]?)(\d{1,2})(?![a-zA-Z0-9])')
RE_SCIENTIFIC_NOTATION = re.compile(r'(10)([-+−])(\d+)')
RE_MULTIPLE_NEWLINES = re.compile(r'\n{3,}')
RE_PURE_NUMBERS = re.compile(r'^[-−]?[\d\s]+$')
RE_SINGLE_DIGIT = re.compile(r'^[1-9]$')


# ========== 工具函数 ==========

def convert_to_subscript(text: str) -> str:
    """将文本转换为下标形式"""
    return ''.join(SCRIPT_MAPS.subscript.get(char, char) for char in text)


def convert_to_superscript(text: str) -> str:
    """将文本转换为上标形式"""
    return ''.join(SCRIPT_MAPS.superscript.get(char, char) for char in text)


def normalize_special_chars(text: str) -> str:
    """规范化特殊字符"""
    for old, new in SPECIAL_CHAR_MAP.items():
        text = text.replace(old, new)
    return text


def is_subscript_context(context: str) -> Tuple[bool, bool]:
    """
    根据上下文判断应该是下标还是上标
    
    Returns:
        Tuple[bool, bool]: (is_subscript, is_superscript)
    """
    if not context:
        return (False, False)
    
    context_stripped = context.rstrip()
    
    # 1. 面积/体积单位后面跟正数 → 上标 (cm², mm³)
    for unit in AREA_VOLUME_UNITS:
        if context_stripped.endswith(unit):
            return (False, True)
    
    # 2. 时间/速率单位后面跟负数 → 上标 (min⁻¹, s⁻¹)
    for unit in RATE_UNITS:
        if context_stripped.endswith(unit):
            return (False, True)
    
    # 3. 科学计数法 10⁻⁵, 10⁺³
    if context_stripped.endswith('10'):
        return (False, True)
    
    # 4. 化学元素后面跟数字 → 下标 (H₂O, Fe₂)
    for elem in CHEMICAL_ELEMENTS:
        if context_stripped.endswith(elem):
            return (True, False)
    
    # 5. 常见化学基团后面跟数字 → 下标 (NH₃, CO₂)
    chem_groups = ['NH', 'NO', 'CO', 'SO', 'PO', 'OH', 'CH', 'FeN', 'TM']
    for group in chem_groups:
        if context_stripped.endswith(group):
            return (True, False)
    
    # 6. 右括号后面跟数字（化学式如 Zn(NO₃)₂）
    if context_stripped.endswith(')'):
        return (True, False)
    
    return (False, False)


# ========== PDF 转换器类 ==========

class PDFToMarkdownConverter:
    """PDF 转 Markdown 转换器"""
    
    def __init__(self, pdf_path: str):
        """初始化转换器"""
        self.pdf_path = pdf_path
        self.pages = []
        self.main_font_size = None
    
    def convert(self) -> Dict[str, Any]:
        """执行转换"""
        import pdfplumber
        
        result = {
            "success": False,
            "text": "",
            "pages": 0,
            "char_count": 0,
            "error": None
        }
        
        try:
            with pdfplumber.open(self.pdf_path) as pdf:
                result["pages"] = len(pdf.pages)
                all_text = []
                
                for i, page in enumerate(pdf.pages):
                    page_text = self._process_page(page, i + 1)
                    all_text.append(page_text)
                
                result["text"] = "\n\n".join(all_text)
                result["char_count"] = len(result["text"])
                result["success"] = True
                
        except Exception as e:
            result["error"] = str(e)
            
        return result
    
    def _process_page(self, page, page_num: int) -> str:
        """处理单页PDF"""
        # 1. 提取表格
        tables = page.find_tables()
        table_bboxes = [t.bbox for t in tables]
        
        # 2. 提取字符信息
        chars = page.chars
        
        if not chars:
            return ""
        
        # 3. 分析主字体大小
        self._analyze_font_sizes(chars)
        
        # 4. 按行组织字符
        lines = self._organize_chars_to_lines(chars, table_bboxes)
        
        # 5. 转换为文本
        text = self._lines_to_text(lines, tables, table_bboxes)
        
        return text
    
    def _analyze_font_sizes(self, chars: List[Dict]) -> None:
        """分析页面中的主要字体大小"""
        sizes = [round(c.get('size', 12), 1) for c in chars if c.get('size')]
        if sizes:
            size_counter = Counter(sizes)
            self.main_font_size = size_counter.most_common(1)[0][0]
    
    def _organize_chars_to_lines(self, chars: List[Dict], table_bboxes: List[Tuple]) -> List[List[Dict]]:
        """将字符组织成行，同时标记下标/上标"""
        if not chars:
            return []
        
        # 找出主字体大小
        sizes = [c.get('size', 12) for c in chars]
        size_counter = Counter([round(s, 1) for s in sizes])
        main_size = size_counter.most_common(1)[0][0] if size_counter else 12
        
        # 过滤掉表格内的字符
        non_table_chars = [
            char for char in chars
            if not self._is_char_in_table(char, table_bboxes)
        ]
        
        # 按位置排序
        sorted_chars = sorted(non_table_chars, key=lambda c: (c.get('top', 0), c.get('x0', 0)))
        
        # 行分组
        lines = self._group_chars_into_lines(sorted_chars, main_size)
        
        # 标记下标/上标
        self._mark_subscripts_superscripts(lines, main_size)
        
        return lines
    
    def _is_char_in_table(self, char: Dict, table_bboxes: List[Tuple]) -> bool:
        """检查字符是否在表格内"""
        char_top = char.get('top', 0)
        char_x0 = char.get('x0', 0)
        for bbox in table_bboxes:
            if bbox[0] <= char_x0 <= bbox[2] and bbox[1] <= char_top <= bbox[3]:
                return True
        return False
    
    def _group_chars_into_lines(self, chars: List[Dict], main_size: float) -> List[List[Dict]]:
        """将字符分组到行中"""
        lines = []
        current_line = []
        current_baseline = None
        
        tolerance = max(2, main_size * 0.2)
        
        for char in chars:
            char_top = char.get('top', 0)
            char_size = char.get('size', 12)
            size_ratio = char_size / main_size if main_size > 0 else 1
            
            # 判断是否是小字体
            is_small_font = size_ratio < 0.85
            
            if current_baseline is None:
                current_line = [char.copy()]
                current_baseline = char_top
            else:
                should_join = False
                
                if is_small_font:
                    # 小字体：检查是否紧接在前一个字符后面
                    if current_line:
                        last_char = current_line[-1]
                        gap = char.get('x0', 0) - last_char.get('x1', 0)
                        max_gap = max(last_char.get('size', 12), char_size) * 0.5
                        vertical_range = main_size * 1.2
                        
                        if gap <= max_gap and abs(char_top - current_baseline) <= vertical_range:
                            should_join = True
                else:
                    # 正常字体：使用标准容差判断
                    if abs(char_top - current_baseline) <= tolerance:
                        should_join = True
                        current_baseline = char_top
                
                if should_join:
                    current_line.append(char.copy())
                else:
                    if current_line:
                        lines.append(current_line)
                    current_line = [char.copy()]
                    current_baseline = char_top
        
        if current_line:
            lines.append(current_line)
        
        return lines
    
    def _mark_subscripts_superscripts(self, lines: List[List[Dict]], main_size: float) -> None:
        """标记下标和上标字符"""
        for line in lines:
            if not line:
                continue
            
            # 计算行的基准线位置
            tops = [c.get('top', 0) for c in line]
            top_counter = Counter([round(t, 0) for t in tops])
            baseline_top = top_counter.most_common(1)[0][0] if top_counter else 0
            
            # 按x坐标排序处理
            sorted_line = sorted(enumerate(line), key=lambda x: x[1].get('x0', 0))
            
            for orig_idx, char in sorted_line:
                char_text = char.get('text', '')
                char_size = char.get('size', 12)
                char_top = char.get('top', 0)
                
                size_ratio = char_size / main_size if main_size > 0 else 1
                
                if size_ratio < 0.85:
                    # 获取前面的上下文
                    context = self._get_context_before(line, orig_idx)
                    
                    # 使用上下文判断
                    is_sub, is_sup = is_subscript_context(context)
                    
                    if not is_sub and not is_sup:
                        # 上下文无法确定，使用位置判断
                        position_diff = char_top - baseline_top
                        
                        if position_diff > 1.0:
                            is_sub, is_sup = True, False
                        elif position_diff < -1.0:
                            is_sub, is_sup = False, True
                        else:
                            # 默认为下标
                            is_sub, is_sup = True, False
                    
                    char['is_subscript'] = is_sub
                    char['is_superscript'] = is_sup
    
    def _get_context_before(self, line: List[Dict], idx: int, max_chars: int = 15) -> str:
        """获取当前字符前面的上下文文本"""
        context = []
        for i in range(max(0, idx - max_chars), idx):
            text = line[i].get('text', '')
            if text and not line[i].get('is_subscript') and not line[i].get('is_superscript'):
                context.append(text)
        return ''.join(context)
    
    def _lines_to_text(self, lines: List[List[Dict]], tables: List, table_bboxes: List[Tuple]) -> str:
        """将行转换为文本"""
        result_parts = []
        
        # 处理非表格区域
        for line in lines:
            line_text = self._process_line(line)
            if line_text.strip():
                result_parts.append(line_text)
        
        # 处理表格
        for table in tables:
            table_text = self._process_table(table)
            if table_text.strip():
                result_parts.append(table_text)
        
        return '\n'.join(result_parts)
    
    def _process_line(self, line: List[Dict]) -> str:
        """处理单行字符"""
        if not line:
            return ""
        
        sorted_chars = sorted(line, key=lambda c: c.get('x0', 0))
        
        result = []
        prev_char = None
        current_subscript = ""
        current_superscript = ""
        
        for char in sorted_chars:
            text = char.get('text', '')
            text = normalize_special_chars(text)
            
            # 处理空格
            if prev_char:
                gap = char.get('x0', 0) - prev_char.get('x1', 0)
                prev_size = prev_char.get('size', 12)
                if gap > prev_size * 0.3:
                    if current_subscript:
                        result.append(convert_to_subscript(current_subscript))
                        current_subscript = ""
                    if current_superscript:
                        result.append(convert_to_superscript(current_superscript))
                        current_superscript = ""
                    result.append(' ')
            
            # 处理下标/上标
            is_sub = char.get('is_subscript', False)
            is_sup = char.get('is_superscript', False)
            
            if is_sub:
                if current_superscript:
                    result.append(convert_to_superscript(current_superscript))
                    current_superscript = ""
                current_subscript += text
            elif is_sup:
                if current_subscript:
                    result.append(convert_to_subscript(current_subscript))
                    current_subscript = ""
                current_superscript += text
            else:
                if current_subscript:
                    result.append(convert_to_subscript(current_subscript))
                    current_subscript = ""
                if current_superscript:
                    result.append(convert_to_superscript(current_superscript))
                    current_superscript = ""
                result.append(text)
            
            prev_char = char
        
        # 处理末尾
        if current_subscript:
            result.append(convert_to_subscript(current_subscript))
        if current_superscript:
            result.append(convert_to_superscript(current_superscript))
        
        return ''.join(result)
    
    def _process_table(self, table) -> str:
        """处理表格"""
        try:
            table_data = table.extract()
            if not table_data or len(table_data) < 2:
                return ""
            
            result = []
            max_cols = max(len(row) for row in table_data)
            
            for i, row in enumerate(table_data):
                cells = []
                for cell in row:
                    cell_text = str(cell) if cell else ''
                    cell_text = normalize_special_chars(cell_text)
                    cell_text = self._post_process_text(cell_text)
                    cells.append(cell_text.strip())
                
                while len(cells) < max_cols:
                    cells.append('')
                
                result.append('| ' + ' | '.join(cells) + ' |')
                
                if i == 0:
                    result.append('| ' + ' | '.join(['---'] * max_cols) + ' |')
            
            return '\n'.join(result)
            
        except Exception:
            return ""
    
    def _post_process_text(self, text: str) -> str:
        """后处理文本"""
        # 处理科学计数法
        text = self._process_scientific_notation(text)
        return text
    
    def _process_scientific_notation(self, text: str) -> str:
        """处理科学计数法"""
        def replace_exp(match):
            sign = match.group(2)
            exp = match.group(3)
            return match.group(1) + convert_to_superscript(sign + exp)
        
        return RE_SCIENTIFIC_NOTATION.sub(replace_exp, text)


# ========== 后处理函数 ==========

# ========== DOCX / DOC 转 Markdown ==========

class DocxToMarkdownConverter:
    """DOCX 转 Markdown 转换器

    使用 python-docx 库提取 DOCX 文件中的段落、标题、表格、列表等
    结构化内容，并转换为 Markdown 格式。
    """

    def __init__(self, docx_path: str):
        """初始化转换器

        Args:
            docx_path: DOCX 文件路径
        """
        self.docx_path = docx_path
        self.paragraph_count = 0
        self.table_count = 0

    @staticmethod
    def _fix_docm_content_type(docx_path: str) -> Optional[str]:
        """修复宏启用文档 (.docm) 的 Content-Type 以便 python-docx 正常解析

        部分 .docx 文件实际上是宏启用格式（Content-Type 为
        application/vnd.ms-word.document.macroEnabled.main+xml），
        python-docx 会拒绝解析。本方法创建一个临时副本，
        将 Content-Type 替换为标准格式后返回临时文件路径。

        Args:
            docx_path: 原始文件路径

        Returns:
            修复后的临时文件路径，如果无需修复则返回 None
        """
        import zipfile

        try:
            with zipfile.ZipFile(docx_path, 'r') as zin:
                ct_data = zin.read('[Content_Types].xml')
                ct_text = ct_data.decode('utf-8')

                # 检查是否包含宏启用的 Content-Type
                macro_ct = 'application/vnd.ms-word.document.macroEnabled.main+xml'
                if macro_ct not in ct_text:
                    return None  # 标准格式，无需修复

                # 替换为标准 Content-Type
                standard_ct = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml'
                ct_text = ct_text.replace(macro_ct, standard_ct)

                # 创建临时副本
                tmp_dir = tempfile.mkdtemp()
                tmp_path = os.path.join(tmp_dir, Path(docx_path).name)

                with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                    for item in zin.infolist():
                        data = zin.read(item.filename)
                        if item.filename == '[Content_Types].xml':
                            data = ct_text.encode('utf-8')
                        zout.writestr(item, data)

                return tmp_path

        except Exception:
            return None

    def convert(self) -> Dict[str, Any]:
        """执行转换

        Returns:
            包含转换结果的字典:
            - success: 是否成功
            - text: Markdown 文本
            - paragraphs: 段落数
            - tables: 表格数
            - char_count: 字符数
            - error: 错误信息
        """
        from docx import Document

        result = {
            "success": False,
            "text": "",
            "paragraphs": 0,
            "tables": 0,
            "char_count": 0,
            "error": None
        }

        tmp_docx_path = None
        try:
            # 尝试直接打开，如果失败则尝试修复宏启用格式
            try:
                doc = Document(self.docx_path)
            except ValueError as e:
                error_msg = str(e)
                if 'macroEnabled' in error_msg or 'not a Word file' in error_msg:
                    # 尝试修复宏启用格式
                    tmp_docx_path = self._fix_docm_content_type(self.docx_path)
                    if tmp_docx_path:
                        doc = Document(tmp_docx_path)
                    else:
                        raise
                else:
                    raise

            markdown_parts = []

            # 遍历文档的 body 元素，按照段落和表格在文档中的实际顺序处理
            from docx.oxml.ns import qn

            body = doc.element.body
            para_idx = 0
            table_idx = 0

            for child in body:
                tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

                if tag == 'p':
                    # 处理段落
                    if para_idx < len(doc.paragraphs):
                        para = doc.paragraphs[para_idx]
                        md_text = self._process_paragraph(para)
                        if md_text is not None:
                            markdown_parts.append(md_text)
                        para_idx += 1

                elif tag == 'tbl':
                    # 处理表格
                    if table_idx < len(doc.tables):
                        table = doc.tables[table_idx]
                        md_table = self._process_table(table)
                        if md_table:
                            markdown_parts.append(md_table)
                        table_idx += 1

            full_text = '\n\n'.join(markdown_parts)
            # 清理多余空行
            full_text = RE_MULTIPLE_NEWLINES.sub('\n\n', full_text).strip()

            result["text"] = full_text
            result["paragraphs"] = para_idx
            result["tables"] = table_idx
            result["char_count"] = len(full_text)
            result["success"] = True

        except ImportError:
            result["error"] = "缺少 python-docx 库，请运行: pip install python-docx"
        except Exception as e:
            result["error"] = str(e)
        finally:
            # 清理临时文件
            if tmp_docx_path:
                try:
                    tmp_dir = os.path.dirname(tmp_docx_path)
                    shutil.rmtree(tmp_dir, ignore_errors=True)
                except Exception:
                    pass

        return result

    def _process_paragraph(self, para) -> Optional[str]:
        """处理单个段落，返回 Markdown 文本

        Args:
            para: python-docx Paragraph 对象

        Returns:
            Markdown 格式的文本，如果是空段落或图片段落则返回 None
        """
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        # 跳过空段落
        if not text:
            return None

        # 根据样式处理标题
        if style_name.startswith('Heading'):
            try:
                level = int(style_name.replace('Heading', '').strip())
                level = min(max(level, 1), 6)  # 限制在 1-6 之间
            except ValueError:
                level = 2
            return f"{'#' * level} {text}"

        # 标题变体样式（中文模板常用）
        heading_styles = {
            'Title': '# ',
            'Subtitle': '## ',
            'Heading 1': '# ',
            'Heading 2': '## ',
            'Heading 3': '### ',
            'Heading 4': '#### ',
            'Heading 5': '##### ',
            'Heading 6': '###### ',
        }
        if style_name in heading_styles:
            return f"{heading_styles[style_name]}{text}"

        # 列表处理
        if style_name.startswith('List'):
            indent_level = self._get_list_indent(para)
            prefix = "  " * indent_level
            # 判断是有序列表还是无序列表
            if text and text[0].isdigit():
                # 尝试保持有序列表格式
                return f"{prefix}{text}"
            else:
                return f"{prefix}- {text}"

        # 引用
        if 'Quote' in style_name or 'quote' in style_name.lower():
            return f"> {text}"

        # 检测加粗/斜体 runs
        formatted_text = self._process_runs(para)

        return formatted_text

    def _process_runs(self, para) -> str:
        """处理段落中的 runs，保留加粗、斜体等格式

        Args:
            para: python-docx Paragraph 对象

        Returns:
            带 Markdown 格式的文本
        """
        parts = []
        for run in para.runs:
            run_text = run.text
            if not run_text:
                continue

            # 特殊字符规范化
            run_text = normalize_special_chars(run_text)

            bold = run.bold
            italic = run.italic

            if bold and italic:
                parts.append(f"***{run_text}***")
            elif bold:
                parts.append(f"**{run_text}**")
            elif italic:
                parts.append(f"*{run_text}*")
            else:
                parts.append(run_text)

        return ''.join(parts) if parts else para.text.strip()

    def _get_list_indent(self, para) -> int:
        """获取列表缩进级别

        Args:
            para: python-docx Paragraph 对象

        Returns:
            缩进级别 (0-based)
        """
        try:
            pPr = para._element.pPr
            if pPr is not None:
                numPr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr')
                if numPr is not None:
                    ilvl = numPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl')
                    if ilvl is not None:
                        return int(ilvl.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '0'))
        except Exception:
            pass
        return 0

    def _process_table(self, table) -> str:
        """处理表格，转换为 Markdown 表格

        Args:
            table: python-docx Table 对象

        Returns:
            Markdown 格式的表格字符串
        """
        try:
            rows_data = []
            for row in table.rows:
                cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip() if cell.text else ''
                    cell_text = normalize_special_chars(cell_text)
                    # 将单元格内的换行替换为空格，避免破坏表格格式
                    cell_text = cell_text.replace('\n', ' ').replace('\r', ' ')
                    cells.append(cell_text)
                rows_data.append(cells)

            if not rows_data:
                return ""

            # 确保所有行列数一致
            max_cols = max(len(row) for row in rows_data)
            for row in rows_data:
                while len(row) < max_cols:
                    row.append('')

            result_lines = []
            for i, row in enumerate(rows_data):
                result_lines.append('| ' + ' | '.join(row) + ' |')
                if i == 0:
                    result_lines.append('| ' + ' | '.join(['---'] * max_cols) + ' |')

            return '\n'.join(result_lines)

        except Exception:
            return ""


def convert_doc_to_docx(doc_path: str, output_dir: Optional[str] = None) -> Dict[str, Any]:
    """将 .doc 文件转换为 .docx 文件

    优先使用 LibreOffice 命令行进行转换，如果不可用则尝试 MS Word COM 自动化（仅 Windows）。

    Args:
        doc_path: .doc 文件的完整路径
        output_dir: 输出目录，默认为系统临时目录

    Returns:
        dict: 包含转换结果:
            - success: 是否成功
            - docx_path: 生成的 .docx 文件路径
            - error: 错误信息
    """
    result = {
        "success": False,
        "docx_path": None,
        "error": None
    }

    doc_file = Path(doc_path)
    if not doc_file.exists():
        result["error"] = f"文件不存在: {doc_path}"
        return result

    if output_dir:
        out_dir = Path(output_dir)
    else:
        out_dir = Path(tempfile.mkdtemp())

    out_dir.mkdir(parents=True, exist_ok=True)

    # 方法1: 尝试使用 LibreOffice
    libreoffice_paths = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "soffice",  # 如果在 PATH 中
    ]

    for lo_path in libreoffice_paths:
        try:
            cmd = [
                lo_path,
                "--headless",
                "--convert-to", "docx",
                "--outdir", str(out_dir),
                str(doc_file)
            ]
            proc = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=120,
                encoding='utf-8',
                errors='replace'
            )
            if proc.returncode == 0:
                expected_docx = out_dir / f"{doc_file.stem}.docx"
                if expected_docx.exists():
                    result["success"] = True
                    result["docx_path"] = str(expected_docx)
                    return result
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
        except Exception:
            continue

    # 方法2: 尝试使用 MS Word COM 自动化（仅 Windows）
    # 优先使用 win32com (pywin32)，回退到 comtypes
    for com_lib in ['win32com.client', 'comtypes.client']:
        try:
            if com_lib == 'win32com.client':
                import win32com.client
                word = win32com.client.Dispatch('Word.Application')
            else:
                import comtypes.client
                word = comtypes.client.CreateObject('Word.Application')

            word.Visible = False

            doc_abs = str(doc_file.absolute())
            docx_abs = str((out_dir / f"{doc_file.stem}.docx").absolute())

            # wdFormatXMLDocument = 12
            doc = word.Documents.Open(doc_abs)
            doc.SaveAs(docx_abs, FileFormat=12)
            doc.Close()
            word.Quit()

            result["success"] = True
            result["docx_path"] = docx_abs
            return result

        except ImportError:
            continue
        except Exception as e:
            result["error"] = f"MS Word COM 转换失败 ({com_lib}): {e}"
            return result

    result["error"] = (
        "无法转换 .doc 文件：未找到 LibreOffice 或 MS Word。\n"
        "请安装以下任一工具:\n"
        "  1. LibreOffice (推荐): https://www.libreoffice.org/download/\n"
        "  2. Microsoft Word 并安装 pywin32: pip install pywin32"
    )
    return result


def _get_unique_output_path(output_dir: str, base_name: str, overwrite: bool = False) -> Path:
    """生成唯一的输出文件路径，如果文件名重复则在末尾添加序号

    Args:
        output_dir: 输出目录路径
        base_name: 基础文件名（不含扩展名）
        overwrite: 是否覆盖已存在的文件

    Returns:
        Path: 唯一的输出文件路径
    """
    out_dir = Path(output_dir)
    candidate = out_dir / f"{base_name}.md"

    if overwrite or not candidate.exists():
        return candidate

    # 文件已存在且不覆盖，添加序号
    seq = 2
    while True:
        candidate = out_dir / f"{base_name}_{seq}.md"
        if not candidate.exists():
            return candidate
        seq += 1


def convert_docx_to_markdown(docx_path: str, output_dir: str, overwrite: bool = False) -> dict:
    """将 DOCX 文件转换为 Markdown 格式

    Args:
        docx_path: DOCX 文件的完整路径
        output_dir: 输出 Markdown 文件的目录路径
        overwrite: 是否覆盖已存在的文件

    Returns:
        dict: 包含转换结果:
            - success: 是否成功
            - input_file: 输入文件路径
            - output_file: 输出 Markdown 文件路径
            - char_count: 字符数
            - paragraphs: 段落数
            - tables: 表格数
            - error: 错误信息
    """
    result = {
        "success": False,
        "input_file": docx_path,
        "output_file": None,
        "char_count": 0,
        "paragraphs": 0,
        "tables": 0,
        "error": None
    }

    try:
        docx_name = Path(docx_path).stem

        converter = DocxToMarkdownConverter(docx_path)
        docx_result = converter.convert()

        if not docx_result["success"]:
            result["error"] = f"提取失败: {docx_result.get('error')}"
            return result

        # 后处理：复用 PDF 的化学式修复和智能格式化
        formatted_text = smart_format_text(docx_result["text"])

        # 使用原始文件名命名，如有重复自动添加序号
        output_file = _get_unique_output_path(output_dir, docx_name, overwrite)

        # 构建 Markdown 内容
        markdown_lines = [
            f"# {docx_name}",
            "",
            f"> **Source**: {docx_path}",
            f"> **Paragraphs**: {docx_result['paragraphs']}",
            f"> **Tables**: {docx_result['tables']}",
            f"> **Characters**: {docx_result['char_count']}",
            f"> **Converted**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "",
            "---",
            "",
            formatted_text,
        ]

        # 写入文件
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write('\n'.join(markdown_lines))

        result["success"] = True
        result["output_file"] = str(output_file)
        result["char_count"] = docx_result["char_count"]
        result["paragraphs"] = docx_result["paragraphs"]
        result["tables"] = docx_result["tables"]

    except Exception as e:
        result["error"] = str(e)

    return result


def convert_doc_to_markdown(doc_path: str, output_dir: str, overwrite: bool = False) -> dict:
    """将 DOC 文件转换为 Markdown 格式

    先将 .doc 转换为 .docx，然后按 .docx 流程处理。

    Args:
        doc_path: DOC 文件的完整路径
        output_dir: 输出 Markdown 文件的目录路径
        overwrite: 是否覆盖已存在的文件

    Returns:
        dict: 包含转换结果（字段同 convert_docx_to_markdown），
              另外可能包含 temp_docx_path（临时 docx 文件路径）
    """
    result = {
        "success": False,
        "input_file": doc_path,
        "output_file": None,
        "char_count": 0,
        "paragraphs": 0,
        "tables": 0,
        "error": None,
        "temp_docx_path": None
    }

    try:
        # Step 1: doc → docx
        # 注意：不在此处预检查输出文件，因为最终文件名取决于提取的论文标题
        temp_dir = tempfile.mkdtemp()
        convert_result = convert_doc_to_docx(doc_path, temp_dir)

        if not convert_result["success"]:
            result["error"] = f"DOC → DOCX 转换失败: {convert_result.get('error')}"
            return result

        docx_path = convert_result["docx_path"]
        result["temp_docx_path"] = docx_path

        # Step 2: docx → markdown
        docx_result = convert_docx_to_markdown(docx_path, output_dir, overwrite)

        # 清理临时文件
        try:
            if temp_dir:
                shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception:
            pass

        # 汇总结果
        result["success"] = docx_result["success"]
        result["output_file"] = docx_result.get("output_file")
        result["char_count"] = docx_result.get("char_count", 0)
        result["paragraphs"] = docx_result.get("paragraphs", 0)
        result["tables"] = docx_result.get("tables", 0)
        result["error"] = docx_result.get("error")

    except Exception as e:
        result["error"] = str(e)

    return result


# ========== 后处理函数 ==========

def post_process_text(text: str) -> str:
    """
    后处理文本 - 修复常见的转换问题
    """
    # 1. 合并被分割的下标/上标行
    text = merge_broken_script_lines(text)
    
    # 2. 修复化学式中的数字下标
    text = fix_chemical_formulas(text)
    
    # 3. 修复单位中的上标
    text = fix_unit_superscripts(text)
    
    # 4. 清理多余的空行
    text = RE_MULTIPLE_NEWLINES.sub('\n\n', text)
    
    return text.strip()


def merge_broken_script_lines(text: str) -> str:
    """
    合并被错误分割的下标/上标行
    
    例如:
    "Fe @NG DAC." + "₂" → "Fe₂@NG DAC."
    "NH  Enabled" + "₂₃" → "NH₃ Enabled"
    """
    lines = text.split('\n')
    merged_lines = []
    i = 0
    
    while i < len(lines):
        current_line = lines[i].rstrip()
        
        # 检查下一行是否是纯下标/上标字符
        if i + 1 < len(lines):
            next_line = lines[i + 1].strip()
            
            # 检测纯下标/上标行（只包含下标或上标字符）
            if next_line and is_pure_script_chars(next_line):
                # 合并到当前行
                # 找到当前行中应该插入下标的位置
                merged_line = insert_script_at_position(current_line, next_line)
                merged_lines.append(merged_line)
                i += 2
                continue
        
        merged_lines.append(current_line)
        i += 1
    
    return '\n'.join(merged_lines)


def is_pure_script_chars(text: str) -> bool:
    """检查文本是否只包含下标或上标字符"""
    if not text:
        return False
    
    # 下标字符范围
    subscript_chars = set('₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₐₑₕᵢⱼₖₗₘₙₒₚᵣₛₜᵤᵥₓ')
    # 上标字符范围
    superscript_chars = set('⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿⁱ')
    
    script_chars = subscript_chars | superscript_chars
    
    # 检查是否所有字符都是下标/上标
    return all(c in script_chars or c.isspace() for c in text)


def insert_script_at_position(line: str, script: str) -> str:
    """
    将下标/上标插入到行中正确的位置
    
    规则:
    1. 如果行以 "@NG" 结尾，下标插入到 @ 前面的元素后
    2. 如果行中有 "NH  rather" 模式，分配下标到空格前的元素
    3. 如果行以元素结尾，插入下标
    4. 否则追加到行末尾
    """
    script = script.strip()
    if not script:
        return line
    
    # 分配多个下标字符
    scripts = list(script)
    script_idx = 0
    
    def get_next_script():
        nonlocal script_idx
        if script_idx < len(scripts):
            s = scripts[script_idx]
            script_idx += 1
            return s
        return ''
    
    # 模式1: "Fe @NG DAC" + "₂" → "Fe₂@NG DAC"
    match = re.search(r'([A-Z][a-z]?)\s+@NG', line)
    if match:
        elem = match.group(1)
        s = get_next_script()
        return line.replace(f'{elem} @NG', f'{elem}{s}@NG', 1)
    
    # 模式2: "NH   rather  than  N" + "₂₃" → "NH₃ rather than N₂"
    # 检测行中有多个空格分隔的元素模式
    match = re.search(r'([A-Z][a-z]?)\s{2,}(rather|than|over|and|or)', line, re.IGNORECASE)
    if match:
        elem1 = match.group(1)
        s1 = get_next_script()
        result = re.sub(rf'{elem1}\s{{2,}}', f'{elem1}{s1} ', line, count=1)
        
        # 检查是否有第二个元素需要分配下标
        match2 = re.search(r'(than|over|and|or)\s+([A-Z][a-z]?)(\s*$|\s+[a-z])', result, re.IGNORECASE)
        if match2 and script_idx < len(scripts):
            elem2 = match2.group(2)
            s2 = get_next_script()
            result = re.sub(rf'(than|over|and|or)\s+{elem2}', rf'\1 {elem2}{s2}', result, count=1, flags=re.IGNORECASE)
        
        return result
    
    # 模式3: "N O" + "₂" → "N₂O"
    match = re.search(r'\b([A-Z][a-z]?)\s+([A-Z][a-z]?)\s*$', line)
    if match:
        elem1 = match.group(1)
        elem2 = match.group(2)
        s = get_next_script()
        return line.replace(f'{elem1} {elem2}', f'{elem1}{s}{elem2}', 1)
    
    # 模式4: 以元素结尾的行
    match = re.search(r'\b([A-Z][a-z]?)\s*$', line)
    if match:
        elem = match.group(1)
        s = get_next_script()
        return line[:match.end(1)] + s + line[match.end(1):]
    
    # 模式5: 行中有 "N synthesis" 这样的模式
    match = re.search(r'\b([A-Z][a-z]?)\s+(synthesis|production|formation|reaction)', line, re.IGNORECASE)
    if match:
        elem = match.group(1)
        s = get_next_script()
        return line.replace(f'{elem} {match.group(2)}', f'{elem}{s} {match.group(2)}', 1)
    
    # 模式6: "N   due" + "₂" → "N₂ due" (元素后面有多个空格)
    match = re.search(r'\b([A-Z][a-z]?)\s{2,}(due|than|over|from|in|on|at)', line, re.IGNORECASE)
    if match:
        elem = match.group(1)
        s = get_next_script()
        return re.sub(rf'\b{elem}\s{{2,}}{match.group(2)}', f'{elem}{s} {match.group(2)}', line, count=1, flags=re.IGNORECASE)
    
    # 模式7: 行以单个元素结尾，后面有多个空格
    match = re.search(r'\b([A-Z][a-z]?)\s{2,}$', line)
    if match:
        elem = match.group(1)
        s = get_next_script()
        return re.sub(rf'{elem}\s{{2,}}$', f'{elem}{s}', line)
    
    # 默认: 追加到末尾
    return line + ''.join(scripts[script_idx:])


def fix_chemical_formulas(text: str) -> str:
    """修复化学式中的数字"""
    def replace_formula(match):
        elem = match.group(1)
        num = match.group(2)
        return elem + convert_to_subscript(num)
    
    # 匹配元素符号后跟数字
    return RE_ELEMENT_NUMBER.sub(replace_formula, text)


def fix_unit_superscripts(text: str) -> str:
    """修复单位中的上标"""
    # 修复面积单位: cm2 → cm², mm2 → mm²
    for unit in AREA_VOLUME_UNITS:
        pattern = rf'({unit})(\d)(?![a-zA-Z0-9])'
        text = re.sub(pattern, lambda m: m.group(1) + convert_to_superscript(m.group(2)), text)
    
    # 修复速率单位负指数: min-1 → min⁻¹
    for unit in RATE_UNITS:
        pattern = rf'({unit})\s*[-−]?(\d)(?![a-zA-Z0-9])'
        text = re.sub(pattern, lambda m: m.group(1) + convert_to_superscript('⁻' + m.group(2) if m.group(0).count('-') or m.group(0).count('−') else m.group(2)), text)
    
    return text


def smart_format_text(text: str) -> str:
    """
    智能格式化文本
    """
    # 后处理
    text = post_process_text(text)
    
    lines = text.split('\n')
    result = []
    
    for line in lines:
        stripped = line.strip()
        
        if not stripped:
            if result and result[-1] != '':
                result.append('')
        else:
            # 检测可能的标题
            if _is_likely_heading(stripped):
                if result and result[-1] != '':
                    result.append('')
                if not stripped.startswith('#'):
                    stripped = '## ' + stripped
                result.append(stripped)
                result.append('')
            else:
                result.append(stripped)
    
    final_text = '\n'.join(result)
    return RE_MULTIPLE_NEWLINES.sub('\n\n', final_text).strip()


def _is_likely_heading(line: str) -> bool:
    """判断是否可能是标题"""
    if not line or line.startswith('#'):
        return False
    
    section_patterns = [
        r'^S-\d+$',
        r'^\d+\.\s+[A-Z]',
        r'^(Abstract|Introduction|Methods?|Results?|Discussion|Conclusion|References)',
        r'^(Experimental|Theoretical|Computational)\s+\w+',
        r'^(Table|Figure|Fig\.)\s+\d+',
        r'^Synthesis\s+of',
        r'^Characterizations?$',
    ]
    
    for pattern in section_patterns:
        if re.match(pattern, line, re.IGNORECASE):
            return True
    
    if len(line) < 50 and line.isupper() and len(line.split()) <= 5:
        return True
    
    return False


def _extract_paper_title_with_llm(text: str, api_key: str, api_base: str, model: str = "qwen-plus") -> str:
    """使用 LLM 从论文文本中提取标题（模仿 ingest_markdown.py 的逻辑）
    
    Args:
        text: Markdown 文本（建议传入前 2000 字符以节省 token）
        api_key: OpenAI 兼容 API 的密钥
        api_base: OpenAI 兼容 API 的基础 URL
        model: 使用的 LLM 模型名称，默认 qwen-plus
        
    Returns:
        str: 提取到的论文标题，失败时返回空字符串
    """
    if not text or not text.strip():
        return ""
    
    try:
        from openai import OpenAI
        
        client = OpenAI(api_key=api_key, base_url=api_base)
        
        # 截取前 2000 字符以节省 token
        head_text = text[:2000]
        
        response = client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "system",
                    "content": (
                        "你是一个学术论文标题提取助手。"
                        "从用户提供的论文文本中提取论文标题。"
                        "只返回标题的纯文本，不要包含任何其他内容（如引号、解释、前缀等）。"
                        "如果找不到明确的论文标题，返回空字符串。"
                    )
                },
                {
                    "role": "user",
                    "content": f"请从以下学术论文文本中提取论文标题：\n\n{head_text}"
                }
            ],
            temperature=0.1,
            max_tokens=200,
        )
        
        title = response.choices[0].message.content.strip() if response.choices else ""
        
        if not title:
            return ""
        
        # 清理可能的多余标记
        title = title.strip('"\'""''`')
        if not title:
            return ""
        
        return title
        
    except Exception as e:
        print(f"[WARN] LLM 提取标题失败: {e}")
        return ""


def extract_paper_title_from_text(text: str) -> str:
    """从转换后的文本中提取论文标题
    
    策略：
    0. 优先使用 LLM 从文本中提取标题（与 ingest_markdown.py 保持一致）
    1. 寻找第一个 Markdown 标题行（# 开头）
    2. 寻找第一个非空、非元数据行中符合标题特征的行
    3. 返回空字符串表示未找到
    
    Args:
        text: 转换后的文本
        
    Returns:
        论文标题字符串，未找到返回空字符串
    """
    if not text:
        return ""
    
    # 策略0: 使用 LLM 提取标题（优先级最高）
    try:
        llm_title = _extract_paper_title_with_llm(
            text,
            api_key=config.openai_api_key,
            api_base=config.openai_api_base
        )
        if llm_title:
            return llm_title
    except Exception:
        pass  # LLM 失败时回退到规则匹配
    
    lines = text.split('\n')
    
    # 策略1: 寻找第一个 Markdown 标题（# 开头，但不是我们生成的元数据标题）
    for line in lines:
        stripped = line.strip()
        # 跳过空行和元数据行
        if not stripped or stripped.startswith('>'):
            continue
        # 跳过分隔线
        if stripped == '---':
            continue
        # 找到第一个 Markdown 标题
        if stripped.startswith('#'):
            title = stripped.lstrip('#').strip()
            # 过滤掉看起来不像论文标题的内容
            if len(title) > 10 and len(title) < 500:
                return title
            break
    
    # 策略2: 寻找前几行中符合标题特征的行（全大写或首字母大写，较长）
    for line in lines[:20]:
        stripped = line.strip()
        if not stripped or stripped.startswith('>') or stripped.startswith('#') or stripped == '---':
            continue
        # 标题通常较长（>15字符）且不是以常见非标题词开头
        non_title_prefixes = ['Table', 'Figure', 'Fig.', 'Abstract', 'Supporting', 'Supplementary']
        if len(stripped) > 15 and not any(stripped.startswith(p) for p in non_title_prefixes):
            # 标题通常不含句号结尾（除非是缩写）
            if not stripped.endswith('.') or 'et al' in stripped:
                return stripped
    
    return ""


def sanitize_filename(title: str) -> str:
    """将论文标题转换为合法的文件名
    
    Args:
        title: 论文标题
        
    Returns:
        合法的文件名字符串
    """
    if not title:
        return ""
    
    # 移除或替换不合法的文件名字符
    # Windows 不允许: \ / : * ? " < > |
    filename = title
    for char in ['\\', '/', ':', '*', '?', '"', '<', '>', '|']:
        filename = filename.replace(char, '')
    
    # 移除首尾空格和点号
    filename = filename.strip().strip('.')
    
    # 压缩连续空格为单个空格
    filename = re.sub(r'\s+', ' ', filename)
    
    # 限制文件名长度（Windows MAX_PATH 考虑，留出目录路径和扩展名的空间）
    if len(filename) > 150:
        filename = filename[:150].strip()
    
    return filename


def convert_pdf_to_markdown(pdf_path: str, output_dir: str, overwrite: bool = False) -> dict:
    """
    将 PDF 文件转换为 Markdown 格式，保留化学式和科学符号格式。
    
    该工具用于从学术论文 PDF 中提取文本内容并转换为结构化的 Markdown 格式。
    支持保留化学式的上下标（如 H₂O、Fe²⁺）、科学计数法以及表格结构。
    
    Use this tool when you need to:
    - Extract text content from scientific paper PDF files
    - Convert PDF documents to Markdown format for further processing
    - Preserve chemical formulas with proper subscript/superscript formatting
    - Process academic papers containing tables and equations
    
    Args:
        pdf_path: PDF 文件的完整路径，必须是有效的 .pdf 文件路径
        output_dir: 输出 Markdown 文件的目录路径，如果不存在会自动创建
        overwrite: 是否覆盖已存在的同名 Markdown 文件，默认为 False（跳过已存在文件）
        
    Returns:
        dict: 包含转换结果的字典，结构如下：
            - success (bool): 转换是否成功
            - input_file (str): 输入的 PDF 文件路径
            - output_file (str): 生成的 Markdown 文件路径
            - char_count (int): 提取的字符总数
            - pages (int): PDF 总页数
            - error (str|None): 错误信息，成功时为 None
    
    Example:
        >>> result = convert_pdf_to_markdown("paper.pdf", "markdown_docs")
        >>> print(result["output_file"])  # "markdown_docs/paper.md"
    """
    result = {
        "success": False,
        "input_file": pdf_path,
        "output_file": None,
        "char_count": 0,
        "pages": 0,
        "error": None
    }
    
    try:
        pdf_name = Path(pdf_path).stem
        
        converter = PDFToMarkdownConverter(pdf_path)
        pdf_result = converter.convert()
        
        if not pdf_result["success"]:
            result["error"] = f"提取失败: {pdf_result.get('error')}"
            return result
        
        # 智能格式化
        formatted_text = smart_format_text(pdf_result["text"])
        
        # 使用原始文件名命名，如有重复自动添加序号
        output_file = _get_unique_output_path(output_dir, pdf_name, overwrite)
        
        # 构建 Markdown 内容
        markdown_lines = [
            f"# {pdf_name}",
            "",
            f"> **Source**: {pdf_path}",
            f"> **Pages**: {pdf_result['pages']}",
            f"> **Characters**: {pdf_result['char_count']}",
            f"> **Converted**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "",
            "---",
            "",
            formatted_text,
        ]
        
        # 写入文件
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write('\n'.join(markdown_lines))
        
        result["success"] = True
        result["output_file"] = str(output_file)
        result["char_count"] = pdf_result["char_count"]
        result["pages"] = pdf_result["pages"]
        
    except Exception as e:
        result["error"] = str(e)
    
    return result


def _dispatch_convert(file_info: Dict[str, Any], output_dir: str, overwrite: bool) -> dict:
    """根据文件类型分发到对应的转换函数

    Args:
        file_info: 包含 path, name, file_type 等信息的字典
        output_dir: 输出目录
        overwrite: 是否覆盖已存在文件

    Returns:
        转换结果字典
    """
    file_path = file_info["path"]
    file_type = file_info.get("file_type", "")

    if file_type == "pdf":
        return convert_pdf_to_markdown(file_path, output_dir, overwrite)
    elif file_type == "docx":
        return convert_docx_to_markdown(file_path, output_dir, overwrite)
    elif file_type == "doc":
        return convert_doc_to_markdown(file_path, output_dir, overwrite)
    else:
        # 根据扩展名自动判断
        ext = Path(file_path).suffix.lower()
        if ext == '.pdf':
            return convert_pdf_to_markdown(file_path, output_dir, overwrite)
        elif ext == '.docx':
            return convert_docx_to_markdown(file_path, output_dir, overwrite)
        elif ext == '.doc':
            return convert_doc_to_markdown(file_path, output_dir, overwrite)
        else:
            return {
                "success": False,
                "input_file": file_path,
                "output_file": None,
                "char_count": 0,
                "error": f"不支持的文件格式: {ext}"
            }


def _detect_single_file_type(file_path: str) -> Optional[str]:
    """检测单个文件的类型

    Args:
        file_path: 文件路径

    Returns:
        文件类型字符串 ("pdf", "docx", "doc") 或 None
    """
    ext = Path(file_path).suffix.lower()
    type_map = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'doc',
    }
    return type_map.get(ext)


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description="文档转 Markdown 工具 - 将 PDF/DOCX/DOC 文件转换为 Markdown 格式",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例:\n"
            "  # 转换单个 PDF 文件\n"
            "  python src/pdf_to_markdown.py --pdf-file paper.pdf -o output\n\n"
            "  # 转换单个 DOCX 文件\n"
            "  python src/pdf_to_markdown.py --docx-file document.docx -o output\n\n"
            "  # 转换单个 DOC 文件（需要 LibreOffice 或 MS Word）\n"
            "  python src/pdf_to_markdown.py --doc-file document.doc -o output\n\n"
            "  # 转换目录中的所有文档（PDF + DOCX + DOC）\n"
            "  python src/pdf_to_markdown.py --input-dir ./docs -o output\n\n"
            "  # 仅转换目录中的 PDF 文件\n"
            "  python src/pdf_to_markdown.py --pdf-dir ./pdfs -o output"
        ),
    )

    parser.add_argument(
        "--pdf-dir", "-d",
        type=str,
        default=None,
        help="PDF 文件目录路径 (默认使用 .env 中的配置)"
    )

    parser.add_argument(
        "--pdf-file", "-f",
        type=str,
        default=None,
        help="单个 PDF 文件路径"
    )

    parser.add_argument(
        "--docx-file",
        type=str,
        default=None,
        help="单个 DOCX 文件路径"
    )

    parser.add_argument(
        "--doc-file",
        type=str,
        default=None,
        help="单个 DOC 文件路径（需要 LibreOffice 或 MS Word）"
    )

    parser.add_argument(
        "--input-dir", "-i",
        type=str,
        default=None,
        help="通用输入目录，自动扫描 PDF、DOCX、DOC 文件"
    )

    parser.add_argument(
        "--output-dir", "-o",
        type=str,
        default=str(Path(__file__).parent.parent / "markdown_docs"),
        help="Markdown 输出目录路径 (默认: pdf_qdrant_mvp/markdown_docs)"
    )

    parser.add_argument(
        "--overwrite", "-w",
        action="store_true",
        help="覆盖已存在的文件"
    )

    args = parser.parse_args()

    # 显示配置信息
    console.print(Panel.fit(
        "[bold cyan]文档转 Markdown 工具[/bold cyan]\n"
        "[dim]支持 PDF / DOCX / DOC 格式，精确保留下标/上标格式[/dim]",
        border_style="cyan"
    ))

    # 创建输出目录
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)
    console.print(f"[green]输出目录已创建/确认: {output_dir}[/green]")

    # ========== 确定要处理的文件列表 ==========
    all_files = []  # 每个元素: {"path": ..., "name": ..., "file_type": ...}

    # 优先级：单文件参数 > 通用 input-dir > pdf-dir
    single_file_args = [
        (args.pdf_file, "pdf"),
        (args.docx_file, "docx"),
        (args.doc_file, "doc"),
    ]

    single_file_found = False
    for file_arg, expected_type in single_file_args:
        if file_arg:
            file_path = Path(file_arg)
            if not file_path.exists():
                console.print(f"[red]文件不存在: {file_arg}[/red]")
                sys.exit(1)

            detected_type = _detect_single_file_type(file_arg)
            if detected_type != expected_type:
                console.print(
                    f"[red]文件扩展名不匹配: 期望 .{expected_type}，"
                    f"实际为 {file_path.suffix}[/red]"
                )
                sys.exit(1)

            all_files.append({
                "path": str(file_path),
                "name": file_path.name,
                "file_type": expected_type,
                "size_kb": file_path.stat().st_size / 1024
            })
            console.print(f"[{expected_type.upper()}] 文件: {file_arg}")
            single_file_found = True
            break  # 只处理第一个指定的单文件

    if not single_file_found:
        if args.input_dir:
            # 通用输入目录：扫描所有支持的格式
            input_path = Path(args.input_dir)
            if not input_path.exists():
                console.print(f"[red]输入目录不存在: {args.input_dir}[/red]")
                sys.exit(1)

            console.print(f"输入目录: {args.input_dir}")
            all_files = get_all_document_files(args.input_dir)

            if not all_files:
                console.print(f"[red]未找到任何支持的文档文件: {args.input_dir}[/red]")
                console.print("[dim]支持的格式: .pdf, .docx, .doc[/dim]")
                sys.exit(1)
        else:
            # 回退到原有的 pdf-dir 逻辑
            if args.pdf_dir:
                config.pdf_dir = args.pdf_dir

            pdf_dir_path = Path(config.pdf_dir)
            if not pdf_dir_path.exists():
                console.print(f"[red]PDF 目录不存在: {config.pdf_dir}[/red]")
                sys.exit(1)

            console.print(f"PDF 目录: {config.pdf_dir}")
            pdf_list = get_pdf_files(config.pdf_dir)

            if not pdf_list:
                console.print(f"[red]未找到 PDF 文件: {config.pdf_dir}[/red]")
                sys.exit(1)

            for f in pdf_list:
                f["file_type"] = "pdf"
            all_files = pdf_list

    # 打印文件汇总
    type_counts = Counter(f["file_type"] for f in all_files)
    console.print(f"输出目录: {args.output_dir}")
    summary_parts = []
    for ft in ["pdf", "docx", "doc"]:
        if type_counts.get(ft, 0) > 0:
            summary_parts.append(f"{type_counts[ft]} 个 {ft.upper()}")
    console.print(f"\n[bold]找到 {'、'.join(summary_parts)}，共 {len(all_files)} 个文件[/bold]")

    # ========== 转换统计 ==========
    stats = {
        "total": len(all_files),
        "success": 0,
        "skipped": 0,
        "failed": 0,
        "total_chars": 0,
        "files": []
    }

    # 使用进度条
    with Progress(console=console) as progress:
        overall_task = progress.add_task(
            "[cyan]转换文档文件...",
            total=len(all_files)
        )

        for i in range(len(all_files)):
            file_info = all_files[i]
            progress.update(overall_task, advance=1)

            file_type_tag = file_info.get("file_type", "???").upper()
            console.print(
                f"\n[{i+1}/{len(all_files)}] "
                f"[{file_type_tag}] 处理: {file_info['name']}"
            )

            result = _dispatch_convert(file_info, args.output_dir, args.overwrite)

            if result["success"]:
                stats["success"] += 1
                stats["total_chars"] += result["char_count"]
                file_stat = {
                    "name": file_info["name"],
                    "output": result["output_file"],
                    "chars": result["char_count"],
                    "type": file_info.get("file_type", "unknown"),
                }
                # PDF 特有字段
                if "pages" in result:
                    file_stat["pages"] = result["pages"]
                # DOCX/DOC 特有字段
                if "paragraphs" in result:
                    file_stat["paragraphs"] = result["paragraphs"]
                if "tables" in result:
                    file_stat["tables"] = result["tables"]
                stats["files"].append(file_stat)

                detail_parts = [f"字符数: {result['char_count']}"]
                if "pages" in result:
                    detail_parts.append(f"页数: {result['pages']}")
                if "paragraphs" in result:
                    detail_parts.append(f"段落: {result['paragraphs']}")
                if "tables" in result:
                    detail_parts.append(f"表格: {result['tables']}")

                console.print(f"  [green]成功: {result['output_file']}[/green]")
                console.print(f"  {', '.join(detail_parts)}")
            elif "已存在" in str(result.get("error", "")):
                stats["skipped"] += 1
                console.print(f"  [yellow]跳过: {result['error']}[/yellow]")
            else:
                stats["failed"] += 1
                console.print(f"  [red]失败: {result.get('error')}[/red]")

    # ========== 显示统计 ==========
    console.print("\n")
    console.print("=" * 60)
    console.print(Panel.fit(
        "[bold green]转换完成统计[/bold green]",
        border_style="green"
    ))

    console.print(f"总文件数: {stats['total']}")
    console.print(f"成功: {stats['success']}")
    console.print(f"跳过: {stats['skipped']}")
    console.print(f"失败: {stats['failed']}")
    console.print(f"总字符数: {stats['total_chars']}")

    if stats["files"]:
        console.print("\n[bold]生成的 Markdown 文件:[/bold]")
        for f in stats["files"]:
            detail_parts = [f"{f['chars']} 字符"]
            if "pages" in f:
                detail_parts.append(f"{f['pages']} 页")
            if "paragraphs" in f:
                detail_parts.append(f"{f['paragraphs']} 段落")
            if "tables" in f:
                detail_parts.append(f"{f['tables']} 表格")
            console.print(f"  - [{f.get('type', '').upper()}] {f['output']}: {', '.join(detail_parts)}")

    console.print("\n")
    console.print(Panel(
        "[bold yellow]下一步操作[/bold yellow]\n\n"
        "Markdown 文件已生成，可以使用 MCP RAG 工具添加到 RAG 系统:\n"
        "  - 使用 add_document 工具添加完整文档\n"
        "  - 使用 search_documents 工具搜索文档\n\n"
        f"输出目录: {output_dir.absolute()}",
        border_style="yellow"
    ))


if __name__ == "__main__":
    main()
